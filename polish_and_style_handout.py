from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import xml.etree.ElementTree as ET
import os
import shutil

SRC_ORIG = r"e:\Documents\Github\RnD\2026-DataScience\Foundations of Data Science.docx"
CLEAN = r"e:\Documents\Github\RnD\2026-DataScience\Lecture101-DataScience_clean.docx"
OUT = r"e:\Documents\Github\RnD\2026-DataScience\Lecture101-DataScience_final.docx"
OUT_PDF = r"e:\Documents\Github\RnD\2026-DataScience\Lecture101-DataScience_final.pdf"

# helper to extract plain paragraphs from original
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def extract_paragraphs(path):
    paras = []
    with zipfile.ZipFile(path) as zin:
        doc_xml = zin.read('word/document.xml')
    root = ET.fromstring(doc_xml)
    for p in root.findall('.//w:p', ns):
        texts = []
        for t in p.findall('.//w:t', ns):
            texts.append(t.text or '')
        text = ''.join(texts).strip()
        if text:
            paras.append(text)
    return paras

orig_paras = extract_paragraphs(SRC_ORIG)

# heuristics to find details
def find_first(keyword_list):
    kl = [k.lower() for k in keyword_list]
    for p in orig_paras:
        low = p.lower()
        for k in kl:
            if k in low:
                return p
    return None

found_objectives = find_first(['objective', 'learning objective'])
found_readings = find_first(['reading', 'textbook', 'references', 'reading list'])
found_assessment = find_first(['assignment', 'assessment', 'midterm', 'final', '%'])
found_contact = find_first(['instructor', 'contact', 'email', 'office'])

# open clean doc
doc = Document(CLEAN)

# set margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# set base styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(11)

# Heading sizes
h2 = doc.styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(14)

# Title style
try:
    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(18)
except Exception:
    pass

# header: add university name and optional logo
section = doc.sections[0]
header = section.header
hdr_p = header.paragraphs[0]
hdr_p.text = 'The Islamia University of Bahawalpur — Department of Information Technology'
hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
hdr_p.runs[0].font.size = Pt(9)

# footer: page number placeholder
footer = section.footer
f_p = footer.paragraphs[0]
f_p.text = 'Lecture101 — Foundations of Data Science    |    Page '
# can't insert dynamic page number easily without field; keep static placeholder
f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
f_p.runs[0].font.size = Pt(9)

# attempt to insert logo if present in 2026-DataScience folder as logo.png or logo.jpg
logo_paths = [r'e:\Documents\Github\RnD\2026-DataScience\logo.png', r'e:\Documents\Github\RnD\2026-DataScience\logo.jpg']
for lp in logo_paths:
    if os.path.exists(lp):
        run = hdr_p.add_run()
        run.add_picture(lp, width=Inches(1.0))
        break

# Replace placeholders in document
# mapping of placeholder text -> replacement
replacements = {}
if found_objectives:
    replacements['Understand core concepts and terminology in Data Science.'] = found_objectives
if found_readings:
    replacements['Primary textbook and selected papers as indicated in class.'] = found_readings
if found_assessment:
    replacements['Assignments: 40%, Midterm: 30%, Final: 30%'] = found_assessment
if found_contact:
    replacements['Instructor: Dr. Mustafa Hameed. Office hours: TBD. Email: example@university.edu'] = found_contact

# perform replacements in paragraphs
for para in doc.paragraphs:
    txt = para.text.strip()
    if not txt:
        continue
    for k, v in replacements.items():
        if k in txt:
            para.text = txt.replace(k, v)

# save final docx
doc.save(OUT)
print('Wrote styled final docx:', OUT)

# convert to PDF via docx2pdf if available
try:
    from docx2pdf import convert
    convert(OUT, OUT_PDF)
    print('Wrote PDF:', OUT_PDF)
except Exception as e:
    print('docx2pdf conversion failed or not available:', e)

