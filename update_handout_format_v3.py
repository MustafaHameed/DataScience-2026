from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

BASE = r"e:\Documents\Github\RnD\~DataScience-2026"
IN_DOCX = os.path.join(BASE, 'Lecture101-DataScience_final_v2.docx')
OUT_DOCX = os.path.join(BASE, 'Lecture101-DataScience_final_v3.docx')
OUT_PDF = os.path.join(BASE, 'Lecture101-DataScience_final_v3.pdf')

# look for image files provided by user
img_candidates = []
for name in os.listdir(BASE):
    if name.lower().endswith(('.png', '.jpg', '.jpeg', '.svg')) and 'diagram' in name.lower():
        img_candidates.append(os.path.join(BASE, name))
# if no diagram with 'diagram' in name, accept any image
if not img_candidates:
    for name in os.listdir(BASE):
        if name.lower().endswith(('.png', '.jpg', '.jpeg', '.svg')):
            img_candidates.append(os.path.join(BASE, name))

PREFERRED_IMG = img_candidates[0] if img_candidates else None
if PREFERRED_IMG:
    print('Found image to insert:', PREFERRED_IMG)
else:
    print('No image found in folder; script will keep placeholder text.')

# Load document
doc = Document(IN_DOCX)

# Helper: find paragraph index by text substring
def find_para_index(substrs):
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        for s in substrs:
            if s in txt:
                return i
    return None

# Replace placeholder diagram paragraph (the one starting with 'Diagram: [' ) with the image if available
placeholder_index = find_para_index(['diagram: [insert', 'diagram:', 'data science diagram'])
if placeholder_index is not None and PREFERRED_IMG:
    # remove that paragraph and insert image paragraph
    # python-docx doesn't support deleting paragraph directly; workaround: clear runs
    p = doc.paragraphs[placeholder_index]
    p.clear()
    run = p.add_run()
    run.add_picture(PREFERRED_IMG, width=Inches(6))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    print('Inserted image into document at paragraph', placeholder_index)
else:
    if placeholder_index is not None:
        print('Placeholder found but no image to insert; leaving text.')
    else:
        print('No placeholder paragraph found; appending diagram heading and image at end.')
        from docx.enum.style import WD_STYLE_TYPE
        h = doc.add_paragraph()
        h.style = 'Heading 2'
        h.add_run('Data Science Diagram').bold = True
        if PREFERRED_IMG:
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_picture(PREFERRED_IMG, width=Inches(6))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph('Diagram: [Please insert provided diagram image named diagram.png/jpg/svg in the same folder]')

# Transform the definitions paragraphs into a bulleted list and italicize keywords
# We expect earlier script added a heading 'Definitions & Related Fields' and then paragraphs like 'Term: definition'
# Find the heading index
def_index = find_para_index(['definitions & related fields', 'definitions & related fields'])
if def_index is not None:
    # scan forward collecting term paragraphs until we hit 'Relevance to Data Science' heading
    terms = []
    i = def_index + 1
    while i < len(doc.paragraphs):
        txt = doc.paragraphs[i].text.strip()
        if not txt:
            i += 1
            continue
        if 'relevance to data science' in txt.lower():
            break
        # expect 'Term: definition'
        if ':' in txt:
            term, defin = txt.split(':', 1)
            terms.append((term.strip(), defin.strip()))
        i += 1
    # remove old paragraphs in that range and insert bulleted list
    # Note: python-docx doesn't support deleting arbitrary paragraphs easily; we'll mark them and recreate document body
    # Simpler approach: append a new bulleted section at the end, and later we can consider in-place edits
    doc.add_paragraph()
    doc.add_paragraph('Definitions & Related Fields (reformatted)', style='Heading 2')
    for term, defin in terms:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(term + ': ')
        run.bold = True
        # italicize a few keywords heuristically: words like 'data', 'model', 'analytics', 'learning', 'financial', 'marketing', 'business', 'health'
        keywords = ['data', 'model', 'analytics', 'learning', 'financial', 'marketing', 'business', 'health', 'machine', 'learning']
        # build definition with italics for keywords
        words = defin.split()
        new_run = None
        for w in words:
            clean = ''.join(ch for ch in w if ch.isalnum()).lower()
            if clean in keywords:
                r = p.add_run(w + ' ')
                r.italic = True
            else:
                p.add_run(w + ' ')
    print('Appended reformatted Definitions section at end of document')
else:
    print('Definitions heading not found; appending a new formatted Definitions section')
    doc.add_paragraph('Definitions & Related Fields', style='Heading 2')
    definitions = [
        ('Business Analytics', 'The practice of using data analysis and quantitative methods to support business decision-making.'),
        ('Health Analytics', 'Use of data analysis and models to improve patient outcomes and public health.'),
        ('Business Intelligence (BI)', 'Dashboards and reporting systems that transform raw data into actionable information.'),
        ('Learning Analytics', 'Analysis of learner data to improve educational outcomes.'),
        ('FinTech', 'Technology-driven financial services leveraging data and machine learning.'),
        ('Digital Marketing', 'Data-driven use of digital channels to reach and engage customers.')
    ]
    for term, defin in definitions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(term + ': ').bold = True
        # italicize 'data' and other keywords
        parts = defin.split()
        for w in parts:
            if any(k in w.lower() for k in ['data','model','learning','financial','marketing','business','health']):
                r = p.add_run(w + ' ')
                r.italic = True
            else:
                p.add_run(w + ' ')

# Add a comparison table: columns: Field | Typical Data Sources | Typical Models | Primary Goal
# Place it after the Relevance heading if present, else at end
rel_index = find_para_index(['relevance to data science'])
if rel_index is not None:
    # append table after relevance heading
    insert_pos = rel_index + 1
else:
    insert_pos = len(doc.paragraphs)

# Create table at end (python-docx easier to append than insert)
rows = 7
cols = 4
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Field'
hdr_cells[1].text = 'Typical Data Sources'
hdr_cells[2].text = 'Typical Models'
hdr_cells[3].text = 'Primary Goal'

entries = [
    ('Business Analytics', 'Sales, ERP, CRM', 'Descriptive stats, Regression', 'Improve operations & decisions'),
    ('Health Analytics', 'EHR, Clinical Trials, Sensors', 'Survival models, Classification', 'Improve patient outcomes'),
    ('Business Intelligence (BI)', 'Transaction logs, Warehouses', 'Aggregations, OLAP', 'Reporting & KPIs'),
    ('Learning Analytics', 'LMS logs, Assessments', 'Clustering, Predictive models', 'Personalize learning'),
    ('FinTech', 'Transaction data, Market feeds', 'Time-series, ML classifiers', 'Risk scoring, fraud detection'),
    ('Digital Marketing', 'Clickstream, CRM, Social', 'Recommendation systems, A/B testing', 'Optimize engagement & ROI')
]
for ent in entries:
    row_cells = table.add_row().cells
    for j, val in enumerate(ent):
        row_cells[j].text = val

print('Added comparison table (appended at document end)')

# Save and convert
doc.save(OUT_DOCX)
print('Saved updated DOCX as', OUT_DOCX)
try:
    from docx2pdf import convert
    convert(OUT_DOCX, OUT_PDF)
    print('Converted to PDF:', OUT_PDF)
except Exception as e:
    print('docx2pdf conversion failed:', e)
