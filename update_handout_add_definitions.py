from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

# Paths
BASE = r"e:\Documents\Github\RnD\~DataScience-2026"
IN_DOCX = os.path.join(BASE, 'Lecture101-DataScience_final.docx')
OUT_DOCX = os.path.join(BASE, 'Lecture101-DataScience_final_v2.docx')
OUT_PDF = os.path.join(BASE, 'Lecture101-DataScience_final_v2.pdf')
DIAGRAM_PNG = os.path.join(BASE, 'data_science_diagram.png')

# Create a simple diagram image if PIL available
try:
    from PIL import Image, ImageDraw, ImageFont
    W, H = 1200, 600
    img = Image.new('RGB', (W, H), color='white')
    d = ImageDraw.Draw(img)
    # Draw title
    try:
        fnt = ImageFont.truetype('arial.ttf', 36)
    except Exception:
        fnt = ImageFont.load_default()
    d.text((W//2 - 220, 20), 'Data Science Overview', font=fnt, fill=(0,0,0))
    # Draw boxes and arrows (very simple)
    box_font = ImageFont.truetype('arial.ttf', 20) if os.path.exists('C:\Windows\Fonts\arial.ttf') else ImageFont.load_default()
    boxes = ['Data', 'Processing', 'Modeling', 'Deployment']
    x = 100
    y = 150
    for b in boxes:
        d.rectangle([x, y, x+220, y+100], outline='black')
        w, h = d.textsize(b, font=box_font)
        d.text((x + 110 - w/2, y + 50 - h/2), b, fill=(0,0,0), font=box_font)
        x += 260
    # arrows
    d.line([320, 200, 360, 200], fill='black', width=3)
    d.line([580, 200, 620, 200], fill='black', width=3)
    d.line([840, 200, 880, 200], fill='black', width=3)
    img.save(DIAGRAM_PNG)
except Exception as e:
    # If PIL not available, leave DIAGRAM_PNG absent
    DIAGRAM_PNG = None

# Load document
doc = Document(IN_DOCX)

# Update base styles: Normal -> Times New Roman 12, headings consistent
styles = doc.styles
if 'Normal' in styles:
    normal = styles['Normal']
    try:
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
    except Exception:
        pass
# Make Heading 2 bold and slightly larger
if 'Heading 2' in styles:
    h2 = styles['Heading 2']
    try:
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
    except Exception:
        pass

# Find insertion point: after the first heading 'Data Science' or after Overview
insert_index = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip().lower()
    if 'weekly topics' in text or 'weekly topics / detailed content' in text:
        insert_index = i + 1
        break
if insert_index is None:
    insert_index = len(doc.paragraphs)

# Build new content paragraphs
def add_heading(text):
    p = doc.add_paragraph()
    p.style = 'Heading 2'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    run.bold = True
    return p

def add_paragraph(text):
    p = doc.add_paragraph(text)
    return p

# Insert diagram heading and image near insertion index
# python-docx doesn't support inserting at arbitrary paragraph index easily; we'll append at end
add_heading('Data Science Diagram')
if DIAGRAM_PNG and os.path.exists(DIAGRAM_PNG):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(DIAGRAM_PNG, width=Inches(6))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
else:
    add_paragraph('Diagram: [Insert a diagram illustrating Data → Processing → Modeling → Deployment]')

add_heading('Definitions & Related Fields')

definitions = [
    ('Business Analytics', 'The practice of using data analysis and quantitative methods to support business decision-making, often focused on descriptive and diagnostic analytics to improve operations and strategy.'),
    ('Health Analytics', 'Application of data analysis, statistical models, and machine learning to health-care data to improve patient outcomes, operational efficiency, and public health decisions.'),
    ('Business Intelligence (BI)', 'Tools and systems that play a key role in the strategic planning process of a corporation, transforming raw data into meaningful and useful information for business analysis.'),
    ('Learning Analytics', 'Measurement, collection, analysis and reporting of data about learners and their contexts, for purposes of understanding and optimizing learning and the environments in which it occurs.'),
    ('FinTech', 'Technology-driven innovation in financial services, using data analytics, machine learning, and software to improve financial activities and customer experiences.'),
    ('Digital Marketing', 'Use of digital channels and data-driven techniques to reach and engage customers, optimize campaigns, and measure marketing performance.')
]

for term, defin in definitions:
    p = doc.add_paragraph()
    r = p.add_run(term + ': ')
    r.bold = True
    r2 = p.add_run(defin)
    # emphasize a few keywords in italics
    # simple heuristic: italicize the first key noun after comma if present
    # (keep it straightforward)

add_heading('Relevance to Data Science')
relevance_texts = {
    'Business Analytics': 'Business analytics overlaps with data science on data cleaning, visualization, and modeling; data scientists provide predictive models while BA focuses on translating insights to business action.',
    'Health Analytics': 'Health analytics leverages data science methods (EHR analysis, predictive risk modeling) to enable evidence-based clinical decisions and population health management.',
    'Business Intelligence (BI)': 'BI provides dashboards and historical reporting; data science complements BI by adding predictive and prescriptive analytics for forward-looking decisions.',
    'Learning Analytics': 'Learning analytics uses data science to personalize learning pathways, detect at-risk students, and measure educational interventions.',
    'FinTech': 'FinTech applies machine learning for credit scoring, fraud detection, and algorithmic trading—areas where data science models are core.',
    'Digital Marketing': 'Digital marketing relies on data science for customer segmentation, recommendation systems, A/B testing, and campaign optimization.'
}

for term, rel in relevance_texts.items():
    p = doc.add_paragraph()
    r = p.add_run(term + ': ')
    r.bold = True
    p.add_run(rel)

# Save new docx
doc.save(OUT_DOCX)
print('Saved updated DOCX to', OUT_DOCX)

# Try to convert to PDF using docx2pdf if available
try:
    from docx2pdf import convert
    convert(OUT_DOCX, OUT_PDF)
    print('Converted to PDF:', OUT_PDF)
except Exception as e:
    print('Conversion to PDF failed:', e)
