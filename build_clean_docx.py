"""
Build a clean DOCX using python-docx API. This script attempts to import python-docx;
if it's not available it will fall back to a simple XML-based write (less styled).
"""
import zipfile
import xml.etree.ElementTree as ET
import os
import sys

SRC = r"e:\Documents\Github\RnD\2026-DataScience\Foundations of Data Science.docx"
OUT = r"e:\Documents\Github\RnD\2026-DataScience\Lecture101-DataScience_clean.docx"

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def get_paragraphs_from_docx(path):
    paras = []
    with zipfile.ZipFile(path) as zin:
        doc_xml = zin.read('word/document.xml')
    root = ET.fromstring(doc_xml)
    for p in root.findall('.//w:p', ns):
        texts = []
        for t in p.findall('.//w:t', ns):
            texts.append(t.text or '')
        text = ''.join(texts).strip()
        if text:
            paras.append(text)
    return paras

paras = get_paragraphs_from_docx(SRC)
if not paras:
    print('No paragraphs extracted; aborting')
    sys.exit(1)

# try python-docx
try:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # Title
    doc.add_heading('Lecture101 — Foundations of Data Science', level=0)
    doc.add_paragraph('')

    # Overview (first two sentences)
    flat = ' '.join(paras)
    sentences = [s.strip() for s in flat.split('.') if s.strip()]
    overview = (sentences[0] + '.') if sentences else 'Overview content.'
    if len(sentences) > 1:
        overview = sentences[0] + '. ' + sentences[1] + '.'
    doc.add_heading('Overview', level=2)
    p = doc.add_paragraph(overview)
    p.style = doc.styles['Normal']

    # Objectives
    doc.add_heading('Learning Objectives', level=2)
    doc.add_paragraph('• Understand core concepts and terminology in Data Science.')
    doc.add_paragraph('• Familiarize with common data analysis workflows and tools.')
    doc.add_paragraph('• Apply basic statistical and machine learning techniques to datasets.')

    # Prereqs
    doc.add_heading('Prerequisites', level=2)
    doc.add_paragraph('Basic programming (Python) and introductory statistics recommended.')

    # Weekly Topics
    doc.add_heading('Weekly Topics / Detailed Content', level=2)
    for ptext in paras:
        doc.add_paragraph(ptext)

    # Readings & Assessment
    doc.add_heading('Readings', level=2)
    doc.add_paragraph('Primary textbook and selected papers as indicated in class.')
    doc.add_heading('Assessment', level=2)
    doc.add_paragraph('Assignments: 40%, Midterm: 30%, Final: 30%')
    doc.add_heading('Contact', level=2)
    doc.add_paragraph('Instructor: Dr. Mustafa Hameed. Office hours: TBD. Email: example@university.edu')

    doc.save(OUT)
    print('Wrote clean DOCX (python-docx):', OUT)
except Exception as e:
    print('python-docx not available or error:', e)
    print('Falling back to a simple XML-based docx copy (plain paragraphs)')
    # fallback: copy original and replace document.xml with simple paragraphs
    root = ET.Element('{' + ns['w'] + '}document')
    root.set('xmlns:w', ns['w'])
    body = ET.SubElement(root, '{%s}body' % ns['w'])
    def add_p(text):
        p = ET.SubElement(body, '{%s}p' % ns['w'])
        r = ET.SubElement(p, '{%s}r' % ns['w'])
        t = ET.SubElement(r, '{%s}t' % ns['w'])
        t.text = text
    add_p('Lecture101 — Foundations of Data Science')
    add_p('')
    add_p('Overview')
    flat = ' '.join(paras)
    add_p((flat.split('.')[:2]))
    add_p('')
    add_p('Weekly Topics / Detailed Content')
    for p in paras:
        add_p(p)
    new_doc_xml = ET.tostring(root, encoding='utf-8', method='xml')
    # create new docx by copying parts
    tmpf = tempfile.NamedTemporaryFile(delete=False)
    tmpname = tmpf.name
    tmpf.close()
    with zipfile.ZipFile(SRC) as zin, zipfile.ZipFile(tmpname, 'w') as zout:
        for item in zin.infolist():
            if item.filename == 'word/document.xml':
                zout.writestr(item, new_doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))
    if os.path.exists(OUT):
        os.remove(OUT)
    shutil.move(tmpname, OUT)
    print('Wrote clean DOCX (fallback):', OUT)
